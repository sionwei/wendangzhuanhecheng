import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
import markdown
import uuid
from werkzeug.utils import secure_filename
import io
import tempfile
from docx2pdf import convert

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.txt', '.md']
        self.documents = {}  # 存储处理后的文档
        self.previews = {}   # 存储预览内容
        self.temp_dir = tempfile.mkdtemp()  # 创建临时目录
    
    def create_element(self, name):
        """创建XML元素"""
        return OxmlElement(name)
    
    def create_attribute(self, element, name, value):
        """创建XML属性"""
        element.set(qn(name), value)
    
    def _set_chinese_font(self, run):
        """设置中文字体"""
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def apply_custom_style(self, paragraph, style_config):
        """应用自定义样式到段落"""
        try:
            # 设置字体
            for run in paragraph.runs:
                if style_config.get('font_family_cn'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), style_config['font_family_cn'])
                if style_config.get('font_family_en'):
                    run.font.name = style_config['font_family_en']
                if style_config.get('font_size'):
                    run.font.size = Pt(float(style_config['font_size']))

            # 设置段落格式
            if style_config.get('alignment'):
                paragraph.alignment = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }[style_config['alignment']]

            # 设置行间距
            if style_config.get('line_spacing'):
                if style_config['line_spacing_rule'] == 'multiple':
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph.paragraph_format.line_spacing = float(style_config['line_spacing'])
                else:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    paragraph.paragraph_format.line_spacing = Pt(float(style_config['line_spacing']))

            # 设置段落间距
            if style_config.get('space_before'):
                paragraph.paragraph_format.space_before = Pt(float(style_config['space_before']))
            if style_config.get('space_after'):
                paragraph.paragraph_format.space_after = Pt(float(style_config['space_after']))

            # 设置缩进
            if style_config.get('first_line_indent'):
                paragraph.paragraph_format.first_line_indent = Cm(float(style_config['first_line_indent']))
            if style_config.get('left_indent'):
                paragraph.paragraph_format.left_indent = Cm(float(style_config['left_indent']))
            if style_config.get('right_indent'):
                paragraph.paragraph_format.right_indent = Cm(float(style_config['right_indent']))

        except Exception as e:
            print(f"应用样式时出错: {str(e)}")
    
    def merge_documents(self, files, layout_type='default', style_config=None):
        """合并多个文档"""
        merged_doc = Document()
        # 设置默认字体
        style = merged_doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc_id = str(uuid.uuid4())
        preview_text = []  # 用于存储预览文本
        
        try:
            for file in files:
                file_ext = os.path.splitext(file.filename)[1].lower()
                preview_text.append(f"\n=== {file.filename} ===\n")
                
                if file_ext == '.docx':
                    text = self._merge_docx(file, merged_doc, layout_type, style_config)
                elif file_ext == '.pdf':
                    text = self._merge_pdf(file, merged_doc, layout_type, style_config)
                elif file_ext == '.txt':
                    text = self._merge_txt(file, merged_doc, layout_type, style_config)
                elif file_ext == '.md':
                    text = self._merge_markdown(file, merged_doc, layout_type, style_config)
                
                preview_text.append(text)
            
            self.documents[doc_id] = merged_doc
            self.previews[doc_id] = '\n'.join(preview_text)
            return doc_id
        except Exception as e:
            raise Exception(f"文档合并失败: {str(e)}")
    
    def _merge_docx(self, file, merged_doc, layout_type, style_config=None):
        """合并Word文档"""
        try:
            doc = Document(file)
            preview_text = []
            
            for paragraph in doc.paragraphs:
                if layout_type == 'original':
                    # 保持原格式
                    p = merged_doc.add_paragraph()
                    self._copy_paragraph_format(paragraph, p)
                    
                    for run in paragraph.runs:
                        new_run = p.add_run(run.text)
                        # 复制字体格式
                        if run.font.name:
                            new_run.font.name = run.font.name
                        else:
                            self._set_chinese_font(new_run)
                        
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.bold:
                            new_run.bold = True
                        if run.italic:
                            new_run.italic = True
                        if run.underline:
                            new_run.underline = True
                        # 复制颜色
                        if run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
                elif layout_type == 'custom' and style_config:
                    # 使用自定义格式
                    p = merged_doc.add_paragraph(paragraph.text)
                    self.apply_custom_style(p, style_config)
                else:
                    # 使用默认格式
                    p = merged_doc.add_paragraph(paragraph.text)
                    self._set_chinese_font(p.runs[0])
                
                preview_text.append(paragraph.text)
            
            return '\n'.join(preview_text)
        except Exception as e:
            raise Exception(f"Word文档处理失败: {str(e)}")
    
    def _merge_pdf(self, file, merged_doc, layout_type, style_config=None):
        """合并PDF文档"""
        try:
            pdf = PdfReader(file)
            text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                text.append(page_text)
                p = merged_doc.add_paragraph(page_text)
                if layout_type == 'custom' and style_config:
                    self.apply_custom_style(p, style_config)
                else:
                    self._set_chinese_font(p.runs[0])
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"PDF文档处理失败: {str(e)}")
    
    def _merge_txt(self, file, merged_doc, layout_type, style_config=None):
        """合并文本文档"""
        try:
            text = file.read().decode('utf-8')
            p = merged_doc.add_paragraph(text)
            if layout_type == 'custom' and style_config:
                self.apply_custom_style(p, style_config)
            else:
                self._set_chinese_font(p.runs[0])
            return text
        except Exception as e:
            raise Exception(f"文本文档处理失败: {str(e)}")
    
    def _merge_markdown(self, file, merged_doc, layout_type, style_config=None):
        """合并Markdown文档"""
        try:
            text = file.read().decode('utf-8')
            html = markdown.markdown(text)
            p = merged_doc.add_paragraph(html)
            if layout_type == 'custom' and style_config:
                self.apply_custom_style(p, style_config)
            else:
                self._set_chinese_font(p.runs[0])
            return text
        except Exception as e:
            raise Exception(f"Markdown文档处理失败: {str(e)}")
    
    def get_document(self, doc_id):
        """获取处理后的文档"""
        return self.documents.get(doc_id)
    
    def get_preview(self, doc_id):
        """获取文档预览"""
        return self.previews.get(doc_id)
    
    def export_document(self, doc_id, format='docx'):
        """导出文档为指定格式"""
        doc = self.get_document(doc_id)
        if not doc:
            raise ValueError("文档不存在")
        
        try:
            if format == 'docx':
                return doc
            elif format == 'pdf':
                # 将Word转换为PDF
                temp_docx = os.path.join(self.temp_dir, f'{doc_id}.docx')
                temp_pdf = os.path.join(self.temp_dir, f'{doc_id}.pdf')
                doc.save(temp_docx)
                convert(temp_docx, temp_pdf)
                
                with open(temp_pdf, 'rb') as f:
                    pdf_content = f.read()
                
                # 清理临时文件
                os.remove(temp_docx)
                os.remove(temp_pdf)
                
                return pdf_content
            elif format == 'txt':
                text = self.get_preview(doc_id)
                return text.encode('utf-8')
            elif format == 'md':
                # 将文档转换为Markdown
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                return '\n\n'.join(text).encode('utf-8')
            else:
                raise ValueError(f"不支持的格式：{format}")
        except Exception as e:
            raise Exception(f"导出失败: {str(e)}")
    
    def _copy_paragraph_format(self, source_paragraph, target_paragraph):
        """复制段落格式"""
        try:
            # 复制对齐方式
            target_paragraph.alignment = source_paragraph.alignment
            
            # 复制缩进
            target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
            target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
            target_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
            
            # 复制行距
            target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
            target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
            target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
        except Exception as e:
            print(f"复制段落格式时出错: {str(e)}")
            # 如果格式复制失败，继续处理而不中断 